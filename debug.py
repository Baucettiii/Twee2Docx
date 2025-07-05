# -*- coding: utf-8 -*-

import re
import docx
import argparse
import glob
import os
import random
import time
import sys
from docx.shared import Pt

# Prova a importare networkx e avvisa l'utente se manca
try:
    import networkx as nx
    from networkx.algorithms import community
except ImportError:
    print("ERRORE: La libreria 'networkx' non è installata.")
    print("Per favore, installala eseguendo il comando: pip install networkx")
    exit()

def parse_twee_file(file_path):
    """
    Analizza un file di testo in formato Twee e estrae i passaggi (capitoli).
    """
    print(f"Inizio analisi del file: {file_path}")
    passages = []
    current_passage = None
    ignoring_metadata_block = False

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                match = re.match(r'^::\s*([^[{]+)', line)
                if match:
                    original_id = match.group(1).strip()
                    if original_id in ["StoryTitle", "StoryData"]:
                        ignoring_metadata_block = True
                        if current_passage:
                            current_passage['content'] = current_passage['content'].strip()
                            passages.append(current_passage)
                            current_passage = None
                        continue
                    
                    ignoring_metadata_block = False
                    if current_passage:
                        current_passage['content'] = current_passage['content'].strip()
                        passages.append(current_passage)

                    current_passage = {
                        'original_id': original_id,
                        'title': original_id,
                        'content': ''
                    }
                elif not ignoring_metadata_block and current_passage:
                    current_passage['content'] += line
        
        if not ignoring_metadata_block and current_passage:
            current_passage['content'] = current_passage['content'].strip()
            passages.append(current_passage)

    except FileNotFoundError:
        print(f"Errore: File non trovato a questo percorso: {file_path}")
        return []
    
    print(f"Analisi completata. Trovati {len(passages)} passaggi validi.")
    return passages

def _calculate_total_cost(id_map, links_graph):
    """Calcola il costo totale (distanza media) e la distanza massima di un layout."""
    total_links, total_distance, max_dist = 0, 0, 0
    violating_links = []
    for source_id, linked_ids in links_graph.items():
        if source_id not in id_map: continue
        for dest_id in linked_ids:
            if dest_id in id_map:
                dist = abs(id_map[source_id] - id_map[dest_id])
                total_distance += dist
                total_links += 1
                if dist > max_dist:
                    max_dist = dist
                violating_links.append({'source': source_id, 'dest': dest_id, 'dist': dist})
    avg_dist = (total_distance / total_links) if total_links > 0 else 0
    return avg_dist, max_dist, violating_links

def _targeted_correction(id_map, links_graph, non_locked_ids, min_dist, max_dist, passes):
    """
    Fase di Correzione Mirata: tenta di risolvere le violazioni di max_dist.
    """
    if not max_dist or not passes:
        return id_map
    
    print(f"\n--- Inizio Fase di Correzione Mirata ({passes} passate) ---")
    
    for p in range(passes):
        avg_dist, current_max_dist, all_links = _calculate_total_cost(id_map, links_graph)
        
        violating_links = [link for link in all_links if link['dist'] > max_dist]
        if not violating_links:
            print(f"  -> Passata {p+1}/{passes}: Nessuna violazione trovata. Correzione completata.          ", end='\n')
            break
            
        print(f"  -> Passata {p+1}/{passes}: Trovate {len(violating_links)} violazioni. Tento di risolverle...", end='\r')
        
        violating_links.sort(key=lambda x: x['dist'], reverse=True)
        
        corrections_made = 0
        for link in violating_links:
            node_to_move = link['source']
            target_node = link['dest']
            
            if node_to_move not in non_locked_ids:
                continue

            current_pos = id_map[node_to_move]
            target_pos = id_map[target_node]
            
            landing_zone_start = max(target_pos - max_dist, min(id_map.values()))
            landing_zone_end = min(target_pos + max_dist, max(id_map.values()))
            
            best_swap_candidate = None
            best_cost_reduction = 0

            reverse_id_map = {v: k for k, v in id_map.items()}
            for swap_candidate_pos in range(landing_zone_start, landing_zone_end + 1):
                if swap_candidate_pos == current_pos: continue
                
                swap_candidate_id = reverse_id_map.get(swap_candidate_pos)
                
                if not swap_candidate_id or swap_candidate_id not in non_locked_ids:
                    continue

                id_map_after_swap = id_map.copy()
                id_map_after_swap[node_to_move], id_map_after_swap[swap_candidate_id] = swap_candidate_pos, current_pos
                
                is_valid = True
                for neighbor in links_graph.get(node_to_move, []):
                    if neighbor in id_map_after_swap and abs(id_map_after_swap[node_to_move] - id_map_after_swap[neighbor]) < min_dist:
                        is_valid = False; break
                if not is_valid: continue
                for neighbor in links_graph.get(swap_candidate_id, []):
                     if neighbor in id_map_after_swap and abs(id_map_after_swap[swap_candidate_id] - id_map_after_swap[neighbor]) < min_dist:
                        is_valid = False; break
                if not is_valid: continue

                avg_after, _, _ = _calculate_total_cost(id_map_after_swap, links_graph)
                cost_reduction = avg_dist - avg_after

                if cost_reduction > best_cost_reduction:
                    best_cost_reduction = cost_reduction
                    best_swap_candidate = swap_candidate_id
            
            if best_swap_candidate:
                pos1, pos2 = id_map[node_to_move], id_map[best_swap_candidate]
                id_map[node_to_move], id_map[best_swap_candidate] = pos2, pos1
                corrections_made += 1

        if corrections_made == 0:
            print(f"  -> Passata {p+1}/{passes}: Nessuna ulteriore correzione possibile.                  ", end='\n')
            break
            
    return id_map

def _refine_layout(id_map, links_graph, non_locked_ids, distance, refinement_steps):
    """
    Fase di rifinitura globale che tenta di migliorare il layout tramite scambi casuali.
    """
    if not refinement_steps or not non_locked_ids or len(non_locked_ids) < 2:
        return id_map

    print(f"\n--- Inizio Fase di Rifinitura Globale ({refinement_steps * 1000} tentativi di scambio) ---")
    
    current_avg_dist, current_max_dist, _ = _calculate_total_cost(id_map, links_graph)
    print(f"Stato iniziale: Distanza media: {current_avg_dist:.2f}, Distanza massima: {current_max_dist}")

    swapped_count = 0
    
    for i in range(refinement_steps * 1000):
        orig_id1, orig_id2 = random.sample(non_locked_ids, 2)
        
        if orig_id1 not in id_map or orig_id2 not in id_map: continue

        new_id1, new_id2 = id_map[orig_id1], id_map[orig_id2]

        cost_before, cost_after = 0, 0
        neighbors1 = links_graph.get(orig_id1, []) + [k for k, v in links_graph.items() if orig_id1 in v]
        neighbors2 = links_graph.get(orig_id2, []) + [k for k, v in links_graph.items() if orig_id2 in v]
        
        for n_id in neighbors1:
            if n_id in id_map: cost_before += abs(new_id1 - id_map[n_id]); cost_after += abs(new_id2 - id_map[n_id])
        for n_id in neighbors2:
            if n_id in id_map: cost_before += abs(new_id2 - id_map[n_id]); cost_after += abs(new_id1 - id_map[n_id])

        if cost_after < cost_before:
            is_valid = True
            for n_id in neighbors1:
                if n_id in id_map and abs(new_id2 - id_map[n_id]) < distance: is_valid = False; break
            if not is_valid: continue
            for n_id in neighbors2:
                if n_id in id_map and abs(new_id1 - id_map[n_id]) < distance: is_valid = False; break
            if not is_valid: continue

            id_map[orig_id1], id_map[orig_id2] = new_id2, new_id1
            swapped_count += 1
        
        if i > 0 and i % 10000 == 0:
            avg, max_d, _ = _calculate_total_cost(id_map, links_graph)
            print(f"  ...Rifinitura {i // 1000}k/{refinement_steps}k: Dist. media: {avg:.2f}, Dist. max: {max_d}   ", end='\r')

    print()
    print(f"--- Fase di Rifinitura Globale completata. Eseguiti {swapped_count} scambi migliorativi. ---")
    return id_map

def _place_passages_in_zone(passages_in_zone, available_ids, distance, id_map, links_graph):
    """Funzione ausiliaria per posizionare un gruppo di passaggi in un blocco di ID."""
    processing_queue = list(passages_in_zone)
    random.shuffle(processing_queue)
    deferred_passages = []

    while processing_queue:
        passage_to_place = processing_queue.pop(0)
        original_id = passage_to_place['original_id']
        placed_neighbor_new_ids = {id_map[dest_id] for dest_id in links_graph.get(original_id, []) if dest_id in id_map}
        for src_id, linked_ids in links_graph.items():
            if original_id in linked_ids and src_id in id_map: placed_neighbor_new_ids.add(id_map[src_id])
        valid_positions = [pid for pid in available_ids if all(abs(pid - nid) >= distance for nid in placed_neighbor_new_ids)]

        if valid_positions:
            cost_map = [{'id': pos, 'cost': sum(abs(pos - nid) for nid in placed_neighbor_new_ids)} for pos in valid_positions]
            best_fit_id = min(cost_map, key=lambda x: x['cost'])['id']
            available_ids.remove(best_fit_id)
            id_map[original_id] = best_fit_id
            if deferred_passages:
                processing_queue.extend(deferred_passages)
                deferred_passages = []
        else:
            deferred_passages.append(passage_to_place)

    forced_placements = 0
    if deferred_passages:
        forced_placements = len(deferred_passages)
        for stuck_passage in deferred_passages:
            stuck_original_id = stuck_passage['original_id']
            if not available_ids: 
                print(f"Errore critico: no ID per '{stuck_original_id}'.")
                continue
            forced_id = available_ids.pop(0)
            id_map[stuck_original_id] = forced_id
    
    return forced_placements

def _order_zones_intelligently(communities, G):
    """Ordina le zone (community) in base alla loro interconnessione."""
    if len(communities) <= 1: return communities
    meta_graph = nx.Graph()
    community_map = {node: i for i, com in enumerate(communities) for node in com}
    for u, v in G.edges():
        if u in community_map and v in community_map:
            c1, c2 = community_map[u], community_map[v]
            if c1 != c2:
                if meta_graph.has_edge(c1, c2): meta_graph[c1][c2]['weight'] += 1
                else: meta_graph.add_edge(c1, c2, weight=1)
    if not meta_graph.edges(): random.shuffle(communities); return communities
    path, start_edge = [], max(meta_graph.edges(data=True), key=lambda x: x[2]['weight'])
    path.extend([start_edge[0], start_edge[1]])
    remaining_nodes = set(meta_graph.nodes()) - set(path)
    while remaining_nodes:
        best_candidate, best_score, insert_at_end = None, -1, True
        for node in remaining_nodes:
            score_start = meta_graph.get_edge_data(node, path[0], default={'weight': 0})['weight']
            score_end = meta_graph.get_edge_data(node, path[-1], default={'weight': 0})['weight']
            if score_start > best_score: best_score, best_candidate, insert_at_end = score_start, node, False
            if score_end > best_score: best_score, best_candidate, insert_at_end = score_end, node, True
        if best_candidate:
            if insert_at_end: path.append(best_candidate)
            else: path.insert(0, best_candidate)
            remaining_nodes.remove(best_candidate)
        else: path.extend(list(remaining_nodes)); break
    return [communities[i] for i in path]

def _perform_one_renumbering_attempt(passages, distance, locked_ids, start_number, optimize):
    """Esegue un singolo tentativo completo di rinumerazione per creare una bozza."""
    id_map, passage_dict = {}, {p['original_id']: p for p in passages}
    available_new_ids = list(range(start_number, start_number + len(passages)))
    links_graph, all_passage_ids = {}, set(p['original_id'] for p in passages)
    G = nx.Graph(); G.add_nodes_from(all_passage_ids)
    for p in passages:
        links_graph[p['original_id']] = re.findall(r'\[\[.*?(\d+).*?\]\]', p['content'])
        for dest_id in links_graph[p['original_id']]:
            if dest_id in all_passage_ids: G.add_edge(p['original_id'], dest_id)
    for locked_id in locked_ids:
        if locked_id in passage_dict and available_new_ids:
            id_map[locked_id] = available_new_ids.pop(0)
            print(f"Capitolo bloccato '{locked_id}' -> assegnato al nuovo ID: {id_map[locked_id]}")
    passages_to_renumber = [p for p in passages if p['original_id'] not in locked_ids]
    
    total_forced_placements = 0
    if optimize:
        subgraph_nodes = [p['original_id'] for p in passages_to_renumber]
        subgraph = G.subgraph(subgraph_nodes)
        communities = list(community.greedy_modularity_communities(subgraph))
        ordered_zones = _order_zones_intelligently(communities, G)
        print(f"Trovate e ordinate {len(ordered_zones)} macro-zone.")
        for i, zone in enumerate(ordered_zones):
            zone_passages = [passage_dict[pid] for pid in zone]
            zone_size, zone_available_ids = len(zone_passages), available_new_ids[:len(zone_passages)]
            available_new_ids = available_new_ids[zone_size:]
            if not zone_available_ids: continue
            print(f"  - Posizionamento Zone: {i+1}/{len(ordered_zones)} completato...", end='\r')
            forced_count = _place_passages_in_zone(zone_passages, zone_available_ids, distance, id_map, links_graph)
            total_forced_placements += forced_count
        print() 
    else:
        total_forced_placements = _place_passages_in_zone(passages_to_renumber, available_new_ids, distance, id_map, links_graph)
    
    if total_forced_placements > 0:
        print(f"  - Avviso: {total_forced_placements} capitoli sono stati forzati a causa di vincoli di distanza minima.")

    return id_map, links_graph

def renumber_passages(passages, distance, locked_ids, start_number, optimize, max_dist, attempts, correction_passes, refinement_steps):
    """Funzione principale che gestisce l'intero processo di rinumerazione."""
    best_result, best_avg_dist = None, float('inf')
    for attempt in range(attempts):
        print(f"\n--- Inizio Tentativo di Generazione Bozza {attempt + 1} di {attempts} ---")
        id_map, links_graph = _perform_one_renumbering_attempt(passages, distance, locked_ids, start_number, optimize)
        current_avg_dist, current_max_dist, _ = _calculate_total_cost(id_map, links_graph)
        print(f"Risultato bozza {attempt + 1}: Distanza media: {current_avg_dist:.2f}, Distanza massima: {current_max_dist}")
        if max_dist is None or current_max_dist <= max_dist:
            print(f"Trovata una bozza valida che rispetta la distanza massima ({max_dist})."); best_result = (id_map, links_graph); break
        if current_avg_dist < best_avg_dist: best_avg_dist, best_result = current_avg_dist, (id_map, links_graph)
        if attempt < attempts - 1: print(f"La distanza massima {current_max_dist} viola la soglia di {max_dist}. Riprovo...")
    if best_result is None: print("ERRORE: Nessuna soluzione generata."); return []
    final_id_map, final_links_graph = best_result
    
    final_id_map = _targeted_correction(final_id_map, final_links_graph, [p['original_id'] for p in passages if p['original_id'] not in locked_ids], distance, max_dist, correction_passes)
    
    final_id_map = _refine_layout(final_id_map, final_links_graph, [p['original_id'] for p in passages if p['original_id'] not in locked_ids], distance, refinement_steps)
    
    final_avg_dist, final_max_dist, _ = _calculate_total_cost(final_id_map, final_links_graph)
    print(f"\n--- Statistica finale (dopo tutte le fasi): Distanza media: {final_avg_dist:.2f}, Distanza massima: {final_max_dist} ---")
    
    print("Aggiornamento dei link nei capitoli...")
    updated_passages = []
    for p in passages:
        new_id = final_id_map.get(p['original_id'])
        if new_id is None: continue
        original_links = re.findall(r'\[\[([^\]]+)\]\]', p['content'])
        p['original_links_text'] = ", ".join(original_links) if original_links else "Nessuno"
        new_content = p['content']
        for link_text in original_links:
            old_link_id_match = re.search(r'(\d+)', link_text)
            if old_link_id_match and old_link_id_match.group(1) in final_id_map:
                old_id, new_id_val = old_link_id_match.group(1), final_id_map[old_link_id_match.group(1)]
                updated_link_text = re.sub(r'\b' + re.escape(old_id) + r'\b', str(new_id_val), link_text)
                new_content = new_content.replace(f'[[{link_text}]]', f'[[{updated_link_text}]]')
        p['new_id'], p['content'] = new_id, new_content
        updated_passages.append(p)
    print("Rinumerazione completata."); return updated_passages

def export_to_docx(passages, output_filename, debug_mode, script_start_time):
    """Esporta i passaggi elaborati in un file .docx e stampa il tempo di esecuzione."""
    print(f"Inizio esportazione nel file DOCX: {output_filename}")
    doc = docx.Document()
    sorted_passages = sorted(passages, key=lambda p: p['new_id'])
    for passage in sorted_passages:
        heading = doc.add_heading(f"Capitolo {passage['new_id']}", level=1)
        heading.paragraph_format.keep_with_next = True; heading.paragraph_format.keep_together = True
        p = doc.add_paragraph(); p.paragraph_format.keep_together = True
        content_parts = re.split(r'(\[\[.*?\]\])', passage['content'])
        for part in content_parts:
            if part.startswith('[['):
                link_text = part.strip('[]')
                number_match = re.search(r'(\d+)', link_text)
                if number_match:
                    number = number_match.group(1)
                    before_number, after_number = link_text.split(number, 1)
                    p.add_run(before_number); p.add_run(number).bold = True; p.add_run(after_number)
                else: p.add_run(link_text)
            else: p.add_run(part)
        if debug_mode:
            debug_p = doc.add_paragraph()
            run = debug_p.add_run(f"(Debug: ID Originale: {passage['original_id']}, Rimandi Originali: [{passage.get('original_links_text', 'Nessuno')}])")
            run.italic = True; run.font.size = Pt(8)
    try:
        doc.save(output_filename)
        print(f"Successo! File '{output_filename}' creato correttamente.")
        # SPOSTATO QUI: Calcolo e stampa del tempo
        script_end_time = time.time()
        execution_time = script_end_time - script_start_time
        print(f"\nTempo di esecuzione totale: {execution_time:.2f} secondi.")
        sys.stdout.flush()
    except Exception as e:
        print(f"Errore durante il salvataggio del file DOCX: {e}")

def get_input_file(cli_arg):
    """Determina il file di input."""
    if cli_arg:
        filename = f"{cli_arg}.twee"
        if not os.path.exists(filename): print(f"Errore: File '{filename}' non trovato."); return None
        return filename
    else:
        print("Nessun file specificato, cerco un file .twee..."); twee_files = glob.glob('*.twee')
        if not twee_files: print("Errore: Nessun file .twee trovato."); return None
        print(f"Trovato file: {twee_files[0]}"); return twee_files[0]

def main():
    """Funzione principale che orchestra l'esecuzione dello script."""
    script_start_time = time.time()
    
    parser = argparse.ArgumentParser(description="Processa un file Twee, lo rinumera e lo esporta in DOCX.", formatter_class=argparse.RawTextHelpFormatter)
    parser.add_argument('--nomefile', type=str, help="Nome del file .twee da processare (senza estensione).")
    parser.add_argument('--inizio', type=int, default=1, help="Numero di partenza per la rinumerazione.\nDefault: 1.")
    parser.add_argument('--distanza', type=int, default=5, help="Distanza MINIMA tra capitoli collegati.\nDefault: 5.")
    parser.add_argument('--distanza-max', type=int, help="Distanza MASSIMA accettabile per un singolo rimando.")
    parser.add_argument('--tentativi', type=int, default=1, help="Numero di 'bozze' iniziali da generare.\nDefault: 1.")
    parser.add_argument('--correzione', type=int, default=0, help="Numero di passate di 'correzione mirata' per risolvere\nle violazioni di --distanza-max. Default: 0.")
    parser.add_argument('--rifinitura', type=int, default=0, help="Migliaia di tentativi di 'lucidatura' casuale del layout finale.\nEs: 100 significa 100,000 tentativi. Default: 0.")
    parser.add_argument('--lock', type=str, default='', help="Lista di ID da bloccare, separati da virgola o spazio.\n(es. '1,21,5' o '1 21 5').")
    parser.add_argument('--ottimizza', action=argparse.BooleanOptionalAction, default=True, help="Attiva (default) o disattiva l'ottimizzazione basata su macro-zone.\nUsa --no-ottimizza per disattivarla.")
    parser.add_argument('--debug', action='store_true', help="Attiva le informazioni di debug nel file DOCX.")
    args = parser.parse_args()
    locked_ids = [item.strip() for item in args.lock.replace(',', ' ').split() if item.strip()]
    input_file = get_input_file(args.nomefile)
    if input_file:
        output_file = os.path.splitext(input_file)[0] + '.docx'
        raw_passages = parse_twee_file(input_file)
        if raw_passages:
            final_passages = renumber_passages(
                passages=raw_passages, distance=args.distanza, locked_ids=locked_ids, 
                start_number=args.inizio, optimize=args.ottimizza, max_dist=args.distanza_max,
                attempts=args.tentativi, correction_passes=args.correzione, refinement_steps=args.rifinitura
            )
            if final_passages:
                # Passa il tempo di inizio alla funzione di esportazione
                export_to_docx(final_passages, output_file, args.debug, script_start_time)

# --- ESECUZIONE PRINCIPALE ---
if __name__ == "__main__":
    main()
