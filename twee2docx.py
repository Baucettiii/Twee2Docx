# -*- coding: utf-8 -*-

import re
import docx
import argparse
import glob
import os
import random
import time
import sys
from docx.shared import Pt

# Prova a importare networkx e avvisa l'utente se manca
try:
    import networkx as nx
    from networkx.algorithms import community
except ImportError:
    print("ERRORE: La libreria 'networkx' non è installata.")
    print("Per favore, installala eseguendo il comando: pip install networkx")
    exit()

def parse_twee_file(file_path):
    """
    Analizza un file di testo in formato Twee e estrae i passaggi (capitoli).
    """
    print(f"Inizio analisi del file: {file_path}")
    passages = []
    current_passage = None
    ignoring_metadata_block = False

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                match = re.match(r'^::\s*([^[{]+)', line)
                if match:
                    original_id = match.group(1).strip()
                    if original_id in ["StoryTitle", "StoryData"]:
                        ignoring_metadata_block = True
                        if current_passage:
                            current_passage['content'] = current_passage['content'].strip()
                            passages.append(current_passage)
                            current_passage = None
                        continue
                    
                    ignoring_metadata_block = False
                    if current_passage:
                        current_passage['content'] = current_passage['content'].strip()
                        passages.append(current_passage)

                    current_passage = {
                        'original_id': original_id,
                        'title': original_id,
                        'content': ''
                    }
                elif not ignoring_metadata_block and current_passage:
                    current_passage['content'] += line
        
        if not ignoring_metadata_block and current_passage:
            current_passage['content'] = current_passage['content'].strip()
            passages.append(current_passage)

    except FileNotFoundError:
        print(f"Errore: File non trovato a questo percorso: {file_path}")
        return []
    
    print(f"Analisi completata. Trovati {len(passages)} passaggi validi.")
    return passages

def _print_stats_report(title, id_map, links):
    """Stampa una tabella formattata con le statistiche del layout."""
    avg, max_d, min_d = _calculate_layout_stats(id_map, links)
    print(f"\n--- {title} ---")
    print(f"{'Statistica':<20} | {'Valore':>10}")
    print("-" * 33)
    print(f"{'Distanza Media':<20} | {avg:>10.2f}")
    print(f"{'Distanza Massima':<20} | {max_d:>10}")
    print(f"{'Distanza Minima':<20} | {min_d:>10}")
    print("-" * 33)

def _calculate_layout_stats(id_map, links):
    """
    Calcola le statistiche complete (media, massima, minima) di un layout.
    """
    if not links:
        return 0, 0, 0
        
    distances = []
    for link in links:
        source_id = link['source']
        dest_id = link['dest']
        if source_id in id_map and dest_id in id_map:
            distances.append(abs(id_map[dest_id] - id_map[source_id]))
            
    if not distances:
        return 0, 0, 0

    avg_dist = sum(distances) / len(distances)
    max_dist = max(distances)
    min_dist = min(distances)
    
    return avg_dist, max_dist, min_dist

def _order_zones_intelligently(communities, G):
    """
    Ordina le zone (community) in base alla loro interconnessione.
    """
    if len(communities) <= 1: 
        return communities
        
    meta_graph = nx.Graph()
    community_map = {node: i for i, com in enumerate(communities) for node in com}
    
    for u, v in G.edges():
        if u in community_map and v in community_map:
            c1, c2 = community_map[u], community_map[v]
            if c1 != c2:
                if meta_graph.has_edge(c1, c2): 
                    meta_graph[c1][c2]['weight'] += 1
                else: 
                    meta_graph.add_edge(c1, c2, weight=1)
    
    if not meta_graph.edges(): 
        random.shuffle(communities)
        return communities
    
    try:
        start_edge = max(meta_graph.edges(data=True), key=lambda x: x[2]['weight'])
        path = [start_edge[0], start_edge[1]]
    except ValueError: 
        random.shuffle(communities)
        return communities

    remaining_nodes = set(meta_graph.nodes()) - set(path)
    while remaining_nodes:
        best_candidate, best_score, insert_at_end = None, -1, True
        for node in remaining_nodes:
            score_start = meta_graph.get_edge_data(node, path[0], default={'weight': 0})['weight']
            score_end = meta_graph.get_edge_data(node, path[-1], default={'weight': 0})['weight']
            if score_start > best_score: 
                best_score, best_candidate, insert_at_end = score_start, node, False
            if score_end > best_score: 
                best_score, best_candidate, insert_at_end = score_end, node, True
        
        if best_candidate:
            if insert_at_end: 
                path.append(best_candidate)
            else: 
                path.insert(0, best_candidate)
            remaining_nodes.remove(best_candidate)
        else: 
            path.extend(list(remaining_nodes))
            break
    
    return [communities[i] for i in path]

def _setup_initial_id_mapping(passages, locked_ids, start_number):
    """
    Configura correttamente la mappatura iniziale degli ID rispettando i blocchi.
    """
    print("Configurazione mappatura ID iniziale...")
    
    id_map = {}
    used_ids = set()
    
    # PRIMA: Assegna gli ID bloccati mantenendo il loro valore originale
    for passage in passages:
        original_id = passage['original_id']
        if original_id in locked_ids:
            try:
                numeric_id = int(original_id)
                id_map[original_id] = numeric_id
                used_ids.add(numeric_id)
                print(f"  ID {original_id} bloccato alla posizione {numeric_id}")
            except ValueError:
                print(f"  Avviso: ID bloccato '{original_id}' non è numerico, sarà ignorato")
    
    # SECONDA: Genera ID disponibili per i passaggi non bloccati
    available_ids = []
    current_id = start_number
    needed_ids = len(passages) - len(id_map)
    
    while len(available_ids) < needed_ids:
        if current_id not in used_ids:
            available_ids.append(current_id)
        current_id += 1
    
    print(f"  Generati {len(available_ids)} ID disponibili per i passaggi non bloccati")
    return id_map, available_ids

def _fix_min_dist_violations(id_map, links, non_locked_ids, min_dist, max_passes):
    """
    Corregge le violazioni di distanza minima con un approccio semplificato e più robusto.
    """
    if not max_passes or min_dist <= 1:
        return id_map

    print(f"\n--- Correzione violazioni distanza minima ({max_passes} passate max) ---")
    
    current_map = id_map.copy()
    
    for pass_num in range(max_passes):
        violations = []
        
        # Trova tutte le violazioni attuali
        for link in links:
            source_id, dest_id = link['source'], link['dest']
            if source_id in current_map and dest_id in current_map:
                distance = abs(current_map[dest_id] - current_map[source_id])
                if distance < min_dist:
                    violations.append({
                        'source': source_id,
                        'dest': dest_id,
                        'distance': distance,
                        'deficit': min_dist - distance
                    })
        
        if not violations:
            print(f"  Passata {pass_num + 1}: Nessuna violazione trovata. Correzione completata.")
            break
        
        print(f"  Passata {pass_num + 1}: Trovate {len(violations)} violazioni")
        
        # Ordina per deficit maggiore (violazioni più gravi prima)
        violations.sort(key=lambda x: x['deficit'], reverse=True)
        
        corrections_made = 0
        
        # Prova a correggere le violazioni più gravi
        for violation in violations[:5]:  # Massimo 5 per passata per evitare cicli infiniti
            source_id = violation['source']
            dest_id = violation['dest']
            
            # Determina quale nodo spostare (preferisci quello non bloccato)
            if source_id in non_locked_ids and dest_id not in non_locked_ids:
                node_to_move = source_id
                anchor_node = dest_id
            elif dest_id in non_locked_ids and source_id not in non_locked_ids:
                node_to_move = dest_id
                anchor_node = source_id
            elif source_id in non_locked_ids and dest_id in non_locked_ids:
                # Entrambi non bloccati, scegli casualmente
                node_to_move = random.choice([source_id, dest_id])
                anchor_node = dest_id if node_to_move == source_id else source_id
            else:
                # Entrambi bloccati, salta
                continue
            
            # Trova altri nodi non bloccati con cui scambiare
            anchor_pos = current_map[anchor_node]
            current_pos = current_map[node_to_move]
            
            # Calcola la posizione target ideale
            if current_pos < anchor_pos:
                target_pos = anchor_pos - min_dist
            else:
                target_pos = anchor_pos + min_dist
            
            # Trova il nodo non bloccato più vicino alla posizione target
            best_candidate = None
            best_distance = float('inf')
            
            for candidate_id in non_locked_ids:
                if candidate_id == node_to_move:
                    continue
                
                candidate_pos = current_map[candidate_id]
                distance_to_target = abs(candidate_pos - target_pos)
                
                if distance_to_target < best_distance:
                    # Verifica che lo scambio non crei nuove violazioni
                    if _is_swap_valid(current_map, links, node_to_move, candidate_id, min_dist):
                        best_candidate = candidate_id
                        best_distance = distance_to_target
            
            # Esegui lo scambio se trovato un candidato valido
            if best_candidate:
                current_map[node_to_move], current_map[best_candidate] = current_map[best_candidate], current_map[node_to_move]
                corrections_made += 1
        
        print(f"    Correzioni effettuate: {corrections_made}")
        
        # Se non sono state fatte correzioni, interrompi
        if corrections_made == 0:
            print("    Nessuna ulteriore correzione possibile.")
            break
    
    return current_map

def _is_swap_valid(id_map, links, node1, node2, min_dist):
    """
    Verifica se lo scambio tra due nodi non crea nuove violazioni di distanza minima.
    """
    # Crea una copia della mappa con lo scambio
    test_map = id_map.copy()
    test_map[node1], test_map[node2] = test_map[node2], test_map[node1]
    
    # Verifica tutti i link che coinvolgono i due nodi
    nodes_to_check = [node1, node2]
    
    for node in nodes_to_check:
        # Trova tutti i vicini di questo nodo
        neighbors = set()
        for link in links:
            if link['source'] == node:
                neighbors.add(link['dest'])
            elif link['dest'] == node:
                neighbors.add(link['source'])
        
        # Verifica che tutti i vicini rispettino la distanza minima
        for neighbor in neighbors:
            if neighbor in test_map:
                distance = abs(test_map[node] - test_map[neighbor])
                if distance < min_dist:
                    return False
    
    return True

def renumber_passages_hybrid(passages, min_dist, locked_ids, start_number, correction_passes):
    """
    Motore di rinumerazione ibrido definitivo: Bozza strategica + Correzione robusta.
    """
    if not passages: 
        return [], {}

    print("\n--- Inizio Ottimizzazione Ibrida ---")
    
    # FASE 1: Analisi Strutturale
    print("Fase 1: Analisi della struttura...")
    links = []
    all_passage_ids = {p['original_id'] for p in passages}
    G = nx.Graph()
    G.add_nodes_from(all_passage_ids)
    
    # Miglioramento del pattern regex per catturare meglio i link
    link_pattern = r'\[\[(?:[^\]]*?)(\d+)(?:[^\]]*?)\]\]'
    
    for p in passages:
        for dest_id in re.findall(link_pattern, p['content']):
            if dest_id in all_passage_ids:
                links.append({'source': p['original_id'], 'dest': dest_id})
                G.add_edge(p['original_id'], dest_id)
    
    print(f"  Trovati {len(links)} link tra i passaggi")
    
    # FASE 2: Setup ID mapping corretto
    initial_id_map, available_new_ids = _setup_initial_id_mapping(passages, locked_ids, start_number)
    
    # FASE 3: Generazione layout iniziale per i non bloccati
    print("Fase 3: Generazione layout iniziale...")
    non_locked_ids = [p['original_id'] for p in passages if p['original_id'] not in locked_ids]
    
    if non_locked_ids:
        subgraph = G.subgraph(non_locked_ids)
        
        try:
            if len(subgraph.nodes) > 1:
                communities = list(community.greedy_modularity_communities(subgraph))
            else:
                communities = [set(subgraph.nodes)] if subgraph.nodes else []
        except:
            # Fallback: ogni nodo è una community separata
            communities = [{node} for node in subgraph.nodes]
        
        ordered_zones = _order_zones_intelligently(communities, G)
        
        # Assegna ID alle zone ordinate
        for zone in ordered_zones:
            zone_nodes = sorted(list(zone), key=lambda x: int(x) if x.isdigit() else 0)
            for node in zone_nodes:
                if node not in initial_id_map and available_new_ids:
                    initial_id_map[node] = available_new_ids.pop(0)
    
    # Assegna ID rimanenti
    for passage in passages:
        if passage['original_id'] not in initial_id_map and available_new_ids:
            initial_id_map[passage['original_id']] = available_new_ids.pop(0)

    initial_stats = _calculate_layout_stats(initial_id_map, links)
    _print_stats_report("Statistiche Layout Iniziale", initial_id_map, links)
    
    # FASE 4: Correzione violazioni
    final_id_map = _fix_min_dist_violations(initial_id_map, links, non_locked_ids, min_dist, correction_passes)
    
    final_stats = _calculate_layout_stats(final_id_map, links)
    _print_stats_report("Statistiche Layout Finale", final_id_map, links)
    
    stats_summary = {"before": initial_stats, "after": final_stats}

    print("Aggiornamento dei link nei capitoli...")
    updated_passages = []
    for p in passages:
        new_id = final_id_map.get(p['original_id'])
        if new_id is None: 
            continue
        
        original_links_text = re.findall(r'\[\[([^\]]+)\]\]', p['content'])
        p['original_links_text'] = ", ".join(original_links_text) if original_links_text else "Nessuno"

        new_content = p['content']
        for link_text in original_links_text:
            old_link_id_match = re.search(r'(\d+)', link_text)
            if old_link_id_match and old_link_id_match.group(1) in final_id_map:
                old_id = old_link_id_match.group(1)
                new_id_val = final_id_map[old_id]
                updated_link_text = re.sub(r'\b' + re.escape(old_id) + r'\b', str(new_id_val), link_text)
                new_content = new_content.replace(f'[[{link_text}]]', f'[[{updated_link_text}]]')
        
        p['new_id'] = new_id
        p['content'] = new_content
        updated_passages.append(p)
        
    print("Rinumerazione completata.")
    return updated_passages, stats_summary

def export_to_docx(passages, output_filename, debug_mode, script_start_time, stats):
    """Esporta i passaggi elaborati in un file .docx e stampa le statistiche finali."""
    print(f"Inizio esportazione nel file DOCX: {output_filename}")
    doc = docx.Document()
    sorted_passages = sorted(passages, key=lambda p: p['new_id'])
    
    for passage in sorted_passages:
        heading = doc.add_heading(f"Capitolo {passage['new_id']}", level=1)
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.keep_together = True
        
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        content_parts = re.split(r'(\[\[.*?\]\])', passage['content'])
        
        for part in content_parts:
            if part.startswith('[['):
                link_text = part.strip('[]')
                number_match = re.search(r'(\d+)', link_text)
                if number_match:
                    number = number_match.group(1)
                    before_number, after_number = link_text.split(number, 1)
                    p.add_run(before_number)
                    p.add_run(number).bold = True
                    p.add_run(after_number)
                else: 
                    p.add_run(link_text)
            else: 
                p.add_run(part)
        
        if debug_mode:
            debug_p = doc.add_paragraph()
            run = debug_p.add_run(f"(Debug: ID Originale: {passage['original_id']}, Rimandi Originali: [{passage.get('original_links_text', 'Nessuno')}])")
            run.italic = True
            run.font.size = Pt(8)
    
    try:
        doc.save(output_filename)
        print(f"Successo! File '{output_filename}' creato correttamente.")
        
        before_avg, before_max, before_min = stats['before']
        after_avg, after_max, after_min = stats['after']
        
        print("\n--- Report Statistiche Ottimizzazione (Finale) ---")
        print(f"{'Statistica':<20} | {'Layout Iniziale':>15} | {'Risultato Finale':>18}")
        print("-" * 60)
        print(f"{'Distanza Media':<20} | {before_avg:>15.2f} | {after_avg:>18.2f}")
        print(f"{'Distanza Massima':<20} | {before_max:>15} | {after_max:>18}")
        print(f"{'Distanza Minima':<20} | {before_min:>15} | {after_min:>18}")
        print("-" * 60)

        script_end_time = time.time()
        execution_time = script_end_time - script_start_time
        print(f"\nTempo di esecuzione totale: {execution_time:.2f} secondi.")
        sys.stdout.flush()
    except Exception as e:
        print(f"Errore durante il salvataggio del file DOCX: {e}")

def get_input_file(cli_arg):
    """Determina il file di input."""
    if cli_arg:
        filename = f"{cli_arg}.twee"
        if not os.path.exists(filename): 
            print(f"Errore: File '{filename}' non trovato.")
            return None
        return filename
    else:
        print("Nessun file specificato, cerco un file .twee...")
        twee_files = glob.glob('*.twee')
        if not twee_files: 
            print("Errore: Nessun file .twee trovato.")
            return None
        print(f"Trovato file: {twee_files[0]}")
        return twee_files[0]

def main():
    """Funzione principale che orchestra l'esecuzione dello script."""
    script_start_time = time.time()
    
    parser = argparse.ArgumentParser(
        description="Processa un file Twee, lo rinumera con un algoritmo ibrido e lo esporta.", 
        formatter_class=argparse.RawTextHelpFormatter
    )
    parser.add_argument('--nomefile', type=str, help="Nome del file .twee da processare (senza estensione).")
    parser.add_argument('--inizio', type=int, default=1, help="Numero di partenza per la rinumerazione.\nDefault: 1.")
    parser.add_argument('--distanza-min', type=int, default=10, help="Distanza MINIMA tra capitoli collegati.\nDefault: 10.")
    parser.add_argument('--correzione', type=int, default=20, help="Numero di passate per risolvere le violazioni di distanza minima.\nDefault: 20.")
    parser.add_argument('--lock', type=str, default='', help="Lista di ID da bloccare, separati da virgola o spazio.\n(es. '1,21,5' o '1 21 5').")
    parser.add_argument('--debug', action='store_true', help="Attiva le informazioni di debug nel file DOCX.")
    
    args = parser.parse_args()
    locked_ids = [item.strip() for item in args.lock.replace(',', ' ').split() if item.strip()]
    input_file = get_input_file(args.nomefile)
    
    if input_file:
        output_file = os.path.splitext(input_file)[0] + '.docx'
        raw_passages = parse_twee_file(input_file)
        if raw_passages:
            final_passages, stats = renumber_passages_hybrid(
                passages=raw_passages, 
                min_dist=args.distanza_min,
                locked_ids=locked_ids, 
                start_number=args.inizio,
                correction_passes=args.correzione
            )
            if final_passages:
                export_to_docx(final_passages, output_file, args.debug, script_start_time, stats)

# --- ESECUZIONE PRINCIPALE ---
if __name__ == "__main__":
    main()
